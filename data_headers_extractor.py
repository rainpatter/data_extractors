# Script to extract following parameters from Word documents containing tables and headings:
# - Confidential business information (CBI) - yes/no, category of information
# - Report title
# - ID number of report
# - CAS number with regex (chemical identifier #)
# - Chemical name

from docx import Document
import os
import re
import csv

# Change to folder with documents to read
path = "/home/georgia/Desktop/test assessments"
os.chdir(path)
fileList = [x for x in os.listdir(path) if ".~lock." not in x]


def get_tables(file):
    document = Document(file)
    text_list = []
    try:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        stripped = paragraph.text.strip()
                        final = stripped.upper()
                        text_list.append(final)
    except IndexError:
        pass

    return text_list


def get_text(file):
    document = Document(file)
    full_text = []
    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        final = stripped.upper()
        full_text.append(final)

    return full_text


def get_cbi(all_text):
    list = ["No details are exempt from publication.".upper(),
            'No details are claimed exempt from publication.'.upper()]
    result = any(x in list for x in all_text)
    if result:
        cbi = "No"
    else:
        cbi = "Yes"

    return cbi


def get_cbi_details(all_text):
    exempt_string1 = "Exempt Information  (Section 75 of the Act)".upper()
    exempt_string2 = "Exempt Information (Section 75 of the Act)".upper()
    if exempt_string1 in all_text:
        index = merged.index(exempt_string1)
        index += 1
        cbi = merged[index].strip("Data items and details exempt from publication include".upper())
        cbi_other = merged[index+1]
        all_cbi = cbi + cbi_other
        replaced = all_cbi.replace('.', ', ')
        cbi_list = replaced.split(', ')
        string = 'VARIATION OF DATA REQUIREMENTS (SECTION 24 OF THE ACT)'
        if string in cbi_list:
            cbi_list.remove(string)
            return cbi_list
        else:
            return cbi_list
    elif exempt_string2 in merged:
        index = merged.index(exempt_string2)
        index += 1
        cbi = merged[index].strip("Data items and details exempt from publication include".upper())
        cbi_other = merged[index + 1]
        all_cbi = cbi + cbi_other
        replaced = all_cbi.replace('.', ', ')
        cbi_list = replaced.split(', ')
        string = 'VARIATION OF DATA REQUIREMENTS (SECTION 24 OF THE ACT)'
        if string in cbi_list:
            cbi_list.remove(string)
            return cbi_list
        else:
            return cbi_list
    else:
        pass


def get_report_title(all_text):
    string = "public report".upper()
    for i, item in enumerate(all_text):
        if string in item:
            title_index = all_text.index(item)
    title = next(x for x in merged[:title_index] if x != '')
    return title


def get_assessment_id(all_text):
    try:
        regex = r'\b[A-Z]{2,5}\/\d{1,5}\b'
        match = re.findall(regex, all_text)
        match = [item.strip() for item in match]  # strip whitespace
        match = list(dict.fromkeys(match))  # remove duplicates
        assessment_ids = ' '.join(match)
        return assessment_ids
    except None:
        pass


def get_all_cas_number(all_text):
    try:
        regex = r"\b[1-9]{1}[0-9]{1,5}-\d{2}-\d\b"
        match = re.findall(regex, all_text)
        match = [item.strip() for item in match]  # strip whitespace
        cas_numbers = list(dict.fromkeys(match))  # remove duplicates
        return cas_numbers
    except None:
        pass


def get_chem_name(all_text):
    if 'chemical name'.upper() in all_text:
        index = merged.index("chemical name".upper())
        index += 1
        if merged[index] != '':
            chem_name = merged[index]
            return chem_name
        else:
            index += 2
            chem_name = merged[index]
            return chem_name
    else:
        return "No chemical name"


def get_cas_number(all_text):
    string = 'CAS number'.upper()
    matching = [x for x in all_text if x.startswith(string)]
    if 'CAS number'.upper() in merged:
        index = merged.index('CAS number'.upper())
        index += 2
        cas_number = merged[index]
        return cas_number
    elif matching:
        new = ' '. join(matching)
        stripped = str(new).strip("CAS NUMBER")
        return stripped
    else:
        return "No CAS number"


dataDict = []

for fileItem in fileList:
    text = get_text(fileItem)
    tables = get_tables(fileItem)
    merged_text = ' '.join(text + tables)
    CBI = get_cbi(merged_text)
    CAS_all = get_all_cas_number(merged_text)
    CAS = get_cas_number(merged_text)
    assessmentID = get_assessment_id(merged_text)
    chemName = get_chem_name(merged_text)
    reportTitle = get_report_title(merged_text)
    print(reportTitle)
    if CBI == "Yes":
        CBI_details = get_cbi_details(merged)
    elif CBI == "No":
        CBI_details = "None"
    dict = {"Assessment ID": assessmentID, "CAS number": CAS, "Associated CAS numbers": CAS_all, "CBI": CBI,
            "CBI details": CBIdetails, "Chemical name": chemName}
    dataDict.append(dict)
    print(dict)

desktop = '/home/georgia/Desktop'
os.chdir(desktop)

csv_columns = ['Assessment ID', 'CAS number', 'Associated CAS numbers', 'CBI', 'CBI details', 'Chemical name']
csv_file = 'report_data.csv'

try:
    with open(csv_file, 'w+') as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=csv_columns)
        writer.writeheader()
        for data in dataDict:
            writer.writerow(data)

except IOError:
    print("I/O Error")
