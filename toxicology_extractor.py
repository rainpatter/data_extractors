# Export toxicology data from documents in file directory and extract to spreadsheet

from docx import Document
import os
import re
import csv


testCats = ['ACUTE TOXICITY – ORAL', 'ACUTE TOXICITY – DERMAL',
            'IRRITATION – EYE',
            'IRRITATION – SKIN (IN VITRO)',
            'IRRITATION – SKIN', 'SKIN SENSITISATION',
            'SKIN SENSITISATION – MOUSE LOCAL LYMPH NODE ASSAY (LLNA)',
            'REPEAT DOSE TOXICITY', 'SKIN SENSITISATION – HUMAN VOLUNTEERS',
            'GENOTOXICITY – BACTERIA', 'GENOTOXICITY – IN VITRO',
            'REPRODUCTIVE TOXICITY – TWO GENERATION STUDY',
            'DEVELOPMENTAL TOXICITY'
            'PHOTOSENSITISATION', 'IRRITATION – SKIN (PHOTOTOXIC POTENTIAL)',
            'CONTINUOUS DERMAL IRRITATION', 'CHROMOSOME ABERRATION TEST – IN VITRO'
            ]


path = "/home/georgia/Desktop/test assessments"
os.chdir(path)
fileList = [x for x in os.listdir(path) if ".~lock." not in x]

csv_columns = ['Assessment ID', 'Test', 'Notified or Analogue', 'Conclusion']
csv_file = 'tox_data.csv'


def get_tables(doc_file):
    document = Document(doc_file)
    text_list = []
    try:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        stripped = paragraph.text.strip()
                        final = stripped.upper()
                        text_list.append(final)
    except IndexError:
        pass

    return text_list


def get_text(doc_file):
    document = Document(doc_file)
    full_text = []
    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        final = stripped.upper()
        full_text.append(final)

    return full_text


def get_clean_tables(file_text):
    document = Document(file_text)
    text_list = []
    try:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_list.append(cell.text)
    except IndexError:
        pass
    new_list = []
    for i in range(len(text_list)):
        if text_list[i] != text_list[i-1]:
            new = text_list[i-1]
            stripped = new.strip()
            final = stripped.upper()
            new_list.append(final)
        while "" in new_list:
            new_list.remove("")

    return new_list


def get_assessment_id(text):
    try:
        regex = r'\b[A-Z]{2,5}\/\d{1,5}\b'
        match = re.findall(regex, text)
        match = [item.strip() for item in match]
        match = list(dict.fromkeys(match))
        assessment_ids = ' '.join(match)
        return assessment_ids
    except None:
        pass


def slice_text(text, categories):
    indexes = [i for i, x in enumerate(text) if x in categories]
    parsed = [x for x in indexes if x > 100]
    i = 0
    tox_list = []
    for _ in parsed:
        try:
            slice = text[parsed[i]: parsed[i + 1]]
            tox_list.append(slice)
            i += 1
        except IndexError:
            pass

    return tox_list


def get_tox(tox_list, slice_list, assessment_id):
    data_dict = []
    while True:
        for tox_item in tox_list:
            for slice in slice_list:
                tox_dict = {"Assessment ID": assessment_id,
                            "Test": "None", "Notified or Analogue": "None", "Conclusion": "None"}
                for i, item in enumerate(slice):
                    if tox_item == item:
                        tox_dict["Test"] = tox_item
                        if slice[i + 1] == "TEST SUBSTANCE":
                            tox_dict["Notified or Analogue"] = slice[i+2]
                            if "CONCLUSION" in slice[i:]:
                                index = slice[i:].index("CONCLUSION")
                                tox_dict["Conclusion"] = slice[index + 1]
                                if tox_dict:
                                    data_dict.append(tox_dict)
        return data_dict


masterList = []
for file in fileList:
    cleaned = get_clean_tables(file)
    sliced = slice_text(cleaned, testCats)
    merged = get_tables(file) + get_text(file)
    merged_text = ' '.join(merged)
    assessmentID = get_assessment_id(merged_text)
    tox = get_tox(testCats, sliced, assessmentID)
    masterList.extend(tox)

os.chdir('/home/georgia/Desktop')
try:
    with open(csv_file, 'w+') as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=csv_columns)
        writer.writeheader()
        for data in masterList:
            writer.writerow(data)

except IOError:
    print("I/O Error")




